from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


TEMPLATE = Path(r"C:/Users/hubo0/Downloads/附件1：案例模板.docx")
OUT = Path("基于大模型Skills的交互式Notebook语音识别课程教学案例.docx")


sections = [
    (
        "案例题目",
        "基于大模型 Skills 的交互式 Notebook 在语音识别课程知识内化中的应用研究",
    ),
    (
        "内容与任务（限500字）",
        "本案例面向民办高校“语音识别”课程中 DNN 声学模型、声学特征、训练评估与解码等抽象知识点，针对学生在 AI 编程普及背景下容易“会运行代码、不会解释原理”的问题，设计基于大模型 Skills 的交互式 Notebook 教学任务。教师端使用 asr-teaching-assistant 生成分层教学资源：课堂投屏可采用 AI 在线助教模式，按“示范一步、暂停提问、补讲追问”组织互动；正式练习采用安全分发模式，学生端只运行实验、开放式回答 checkpoint 并导出记录，教师端读取记录后按 rubric 生成学习反馈报告。学生需完成 PyTorch DNN 最小实验，解释声学特征输入、输入/输出维度、训练测试划分、CrossEntropyLoss 与 logits、训练循环、评估模式及 ASR 序列评价等关键问题。任务设置基础理解、代码运行、原理解释、迁移反思三个层次，兼顾课堂互动与个体学习诊断。",
    ),
    (
        "相关知识及背景（限150字）",
        "涉及语音识别基本流程、声学特征、DNN 声学模型、PyTorch 训练循环、模型评估与教育智能体设计。背景来自广东民办高校推进 AI 赋能教学、产教融合与特色发展的需要。",
    ),
    (
        "教学目的（限100字）",
        "引导学生从“运行 AI 生成代码”转向“解释模型原理”，提升语音识别知识内化、工程实践、反思表达与 AI 协同学习能力。",
    ),
    (
        "原理及方案",
        "本案例以大模型 Skills 作为教学流程编排器，将语音识别章节拆解为可运行实验、关键 checkpoint、学生回答记录、教师端 rubric 诊断和个性化反馈五个环节。其基本原理是通过“生成式 AI + 交互式 Notebook + 过程性评价”构建学习闭环：\n"
        "1. 教师端设计 asr-teaching-assistant Skill，内置语音识别课程技术栈、checkpoint 题库、Notebook 模板和反馈报告模板。\n"
        "2. 课堂讲授时，AI 在线助教模式用于投屏示范。AI 每运行一个关键步骤即暂停提问，教师组织学生回答并即时补讲。\n"
        "3. 学生练习时，采用安全分发模式。学生端 Notebook 不含标准答案或判分规则，只记录开放式回答并导出 JSON。\n"
        "4. 教师端 Notebook 读取学生记录，依据 rubric 对“已掌握、部分掌握、需补讲、需教师介入”进行诊断，生成学习反馈报告。\n"
        "5. 对未掌握知识点形成个性化补学建议，服务后续课堂补讲、分层作业和课程持续改进。\n"
        "技术实现上，DNN 示例使用 PyTorch 构建前馈声学模型，用合成声学特征模拟数字识别任务，突出输入特征、输出类别、logits、损失函数、训练循环与评估之间的对应关系。",
    ),
    (
        "报告要求",
        "学生提交内容包括：\n"
        "1. 学生端 Notebook 运行结果截图或导出的运行记录；\n"
        "2. checkpoint 开放式回答记录，需用自己的话解释每一步“为什么这样做”；\n"
        "3. PyTorch DNN 实验结果，包括测试集准确率、混淆矩阵和错误样本观察；\n"
        "4. 对 DNN 声学模型与完整 ASR 系统差异的说明，特别是固定类别分类与 CER/WER 序列评价的区别；\n"
        "5. 课后反思：列出本人最不清楚的 1—2 个知识点，以及准备如何补学。\n"
        "教师端形成学习反馈报告，汇总学生在各 checkpoint 的掌握情况、共性薄弱点和后续教学调整建议。",
    ),
    (
        "考核要求与方法（限300字）",
        "考核采用过程性评价与结果性评价结合。课中依据 checkpoint 回答质量进行即时观察，不以是否写出标准答案为唯一依据，重点考察学生能否解释代码背后的语音识别原理。课后收集学生端 JSON 记录，由教师端 rubric 诊断为“已掌握、部分掌握、需补讲、需教师介入”。评分建议：Notebook 完成度 30%，关键知识点解释 40%，实验结果与错误分析 20%，反思与改进计划 10%。对使用 AI 生成内容的学生，要求能够现场口头说明关键步骤，防止简单复制。教师根据班级薄弱点安排二次补讲或分层练习。",
    ),
    (
        "案例特色或创新（可空缺，限150字）",
        "创新点在于将大模型 Skills 从“代码生成器”转化为“教学流程生成器”，通过学生端/教师端分离降低答案泄露风险，实现课堂互动、个体诊断和个性化补学闭环。",
    ),
    (
        "案例应用（可空缺，限150字）",
        "案例可应用于民办高校人工智能、智能科学、软件工程等专业的语音识别、深度学习实践课程，也可推广到机器学习、自然语言处理、计算机视觉等 AI 课程。",
    ),
]


def set_cell_shading(paragraph, color="EAF2FF"):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    p_pr.append(shd)


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    set_cell_shading(p)
    run = p.add_run(text)
    set_run_font(run, size=12, bold=True, color=(31, 78, 121))
    return p


def add_body(doc, text):
    for idx, part in enumerate(text.split("\n")):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Pt(22) if not part[:2].isdigit() else None
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(part)
        set_run_font(run, size=10.5)
    return p


def main():
    # Load the template to preserve its base section/page setup, then rebuild body.
    doc = Document(TEMPLATE)
    body = doc._body._element
    for child in list(body):
        body.remove(child)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("附件1 案例模板")
    set_run_font(r, size=16, bold=True, color=(31, 78, 121))
    title.paragraph_format.space_after = Pt(14)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = note.add_run("广东民办高校 AI 赋能优秀教学案例")
    set_run_font(rr, size=11, color=(89, 89, 89))
    note.paragraph_format.space_after = Pt(12)

    for heading, content in sections:
        add_heading(doc, heading)
        add_body(doc, content)

    for p in doc.paragraphs:
        if p.alignment is None:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for section in doc.sections:
        section.top_margin = Pt(54)
        section.bottom_margin = Pt(54)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
